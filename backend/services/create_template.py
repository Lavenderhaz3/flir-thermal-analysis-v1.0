"""Create and manage the Word report template."""

import os
from docx import Document
from docx.shared import Pt, Inches


def create_default_template(output_path: str):
    """Create a default docxtpl template with Jinja2 placeholders."""
    doc = Document()

    doc.add_heading("FLIR 红外测温分析报告", level=0)

    doc.add_paragraph("项目名称: {{ project_name }}")
    doc.add_paragraph("报告日期: {{ report_date }}")
    doc.add_paragraph("图片总数: {{ image_count }}")
    doc.add_paragraph("正常设备温度: {{ normal_temp }}°C")

    doc.add_heading("设备测温明细", level=1)

    # Jinja2 for-loop in template
    doc.add_paragraph("{% for img in images %}")
    doc.add_heading("{{ img.filename }}", level=2)

    info_table = doc.add_table(rows=7, cols=2, style="Table Grid")
    rows_data = [
        ("日期", "{{ img.date }}"),
        ("区域", "{{ img.area }}"),
        ("设备编号", "{{ img.equipment }}"),
        ("环境温度 (°C)", "{{ img.ambient_temp }}"),
        ("最高温 (°C)", "{{ img.t_max }}"),
        ("最低温 (°C)", "{{ img.t_min }}"),
        ("平均温 (°C)", "{{ img.t_mean }}"),
    ]
    for i, (label, value) in enumerate(rows_data):
        info_table.rows[i].cells[0].text = label
        info_table.rows[i].cells[1].text = value

    doc.add_paragraph("")
    doc.add_paragraph("框选标注与分析:")

    doc.add_paragraph("{% for ann in img.annotations %}")
    ann_table = doc.add_table(rows=5, cols=2, style="Table Grid")
    ann_data = [
        ("框坐标", "{{ ann.box }}"),
        ("最高温 (°C)", "{{ ann.t_max }}"),
        ("平均温 (°C)", "{{ ann.t_mean }}"),
        ("相对温差", "{{ ann.relative_delta_pct }}"),
        ("状态", "{{ ann.status }}"),
    ]
    for i, (label, value) in enumerate(ann_data):
        ann_table.rows[i].cells[0].text = label
        ann_table.rows[i].cells[1].text = value
    doc.add_paragraph("{% endfor %}")

    doc.add_paragraph("")
    doc.add_paragraph("{% endfor %}")

    doc.add_paragraph("")
    doc.add_paragraph("公式说明: 相对温差 = (最高温 - 正常温度) / (最高温 - 环境温度) × 100%")

    doc.save(output_path)
    return output_path


if __name__ == "__main__":
    out = os.path.join(os.path.dirname(__file__), "..", "templates", "report.docx")
    create_default_template(out)
    print(f"Template saved to {out}")
